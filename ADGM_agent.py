import os
import json
import tempfile
from pathlib import Path
from typing import List, Dict, Tuple

import gradio as gr
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# PDF / DOCX text extraction
import PyPDF2
from docx import Document as DocxReader

# embeddings & vectorstore
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np

# OpenAI for LLM calls
import openai

# CONFIG - Edit as needed

ADGM_REFS_DIR = Path("./adgm_refs")
EMBEDDING_MODEL_NAME = "sentence-transformers/all-mpnet-base-v2"
OPENAI_MODEL = "gpt-4o-mini"   # change if unavailable
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

CHUNK_SIZE = 800
CHUNK_OVERLAP = 150
TOP_K = 3

# Minimal doc-type keyword map (extend this)
DOC_TYPE_KEYWORDS = {
    "Articles of Association": ["articles of association", "aoa"],
    "Memorandum of Association": ["memorandum of association", "moa"],
    "Board Resolution": ["board resolution"],
    "Shareholder Resolution": ["shareholder resolution"],
    "Incorporation Application": ["incorporation application", "application for incorporation"],
    "UBO Declaration": ["ultimate beneficial owner", "ubo declaration"],
    "Register of Members and Directors": ["register of members", "register of directors"],
    "Change of Registered Address": ["change of registered address"]
}

# Checklist mapping for processes; required doc types (simple list)
PROCESS_CHECKLISTS = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Incorporation Application",
        "UBO Declaration",
        "Register of Members and Directors"
    ],
}

# Helpers: text extraction

def extract_text_from_pdf(path: Path) -> str:
    text = []
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for p in reader.pages:
            text.append(p.extract_text() or "")
    return "\n".join(text)

def extract_text_from_docx(path: Path) -> str:
    doc = DocxReader(path)
    return "\n".join([p.text for p in doc.paragraphs])

def load_ref_texts(ref_dir: Path) -> List[str]:
    texts = []
    if not ref_dir.exists():
        return texts
    for file in sorted(ref_dir.iterdir()):
        if file.suffix.lower() == ".pdf":
            texts.append(extract_text_from_pdf(file))
        elif file.suffix.lower() in (".docx", ".doc"):
            texts.append(extract_text_from_docx(file))
        elif file.suffix.lower() == ".txt":
            texts.append(file.read_text(encoding="utf-8"))
    return texts

# Chunking + embeddings + FAISS index

def chunk_text(text: str, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP) -> List[str]:
    chunks = []
    i = 0
    while i < len(text):
        chunk = text[i:i + chunk_size]
        chunks.append(chunk.strip())
        i += chunk_size - overlap
    return [c for c in chunks if len(c) > 50]

class RAGIndex:
    def __init__(self, model_name=EMBEDDING_MODEL_NAME):
        self.model = SentenceTransformer(model_name)
        self.index = None
        self.texts = []

    def build(self, docs: List[str]):
        all_chunks = []
        for doc in docs:
            all_chunks.extend(chunk_text(doc))
        self.texts = all_chunks
        embeddings = self.model.encode(all_chunks, show_progress_bar=True, convert_to_numpy=True)
        dim = embeddings.shape[1]
        self.index = faiss.IndexFlatL2(dim)
        self.index.add(embeddings)
        self.embeddings = embeddings

    def retrieve(self, query: str, top_k=TOP_K) -> List[Tuple[str, float]]:
        q_emb = self.model.encode([query], convert_to_numpy=True)
        D, I = self.index.search(q_emb, top_k)
        results = []
        for idx, dist in zip(I[0], D[0]):
            results.append((self.texts[int(idx)], float(dist)))
        return results

# LLM wrapper (OpenAI)

def openai_analyze(paragraph_text: str, retrieved_chunks: List[str]) -> Dict:
    if OPENAI_API_KEY is None:
        raise RuntimeError("Set OPENAI_API_KEY environment variable for OpenAI usage.")
    openai.api_key = OPENAI_API_KEY

    system = (
        "You are an ADGM compliance assistant. Use ONLY the provided ADGM reference chunks to"
        " evaluate the paragraph. Return JSON ONLY with keys: flag (bool), issues (list), confidence (float)."
    )
    prompt_user = f"""
ADGM REFERENCE CHUNKS:
{"---\n".join(retrieved_chunks)}

PARAGRAPH:
\"\"\"{paragraph_text}\"\"\"

TASK:
1) Based only on the ADGM REFERENCE CHUNKS above, answer whether the paragraph has a compliance issue (flag true/false).
2) If flag is true, issues[] should include objects with keys:
   - reason: short explanation
   - recommendation: suggested edit or clause (1-2 sentences)
   - reference: copy of the relevant ADGM chunk text (or short excerpt)
Return EXACT JSON.
"""

    response = openai.ChatCompletion.create(
        model=OPENAI_MODEL,
        messages=[{"role": "system", "content": system}, {"role": "user", "content": prompt_user}],
        temperature=0.0,
        max_tokens=512,
    )
    raw = response["choices"][0]["message"]["content"].strip()
    try:
        parsed = json.loads(raw)
    except Exception:
        parsed = {"flag": False, "issues": [], "confidence": 0.0, "raw_output": raw}
    return parsed

# Doc type detection & checklist verification
def detect_doc_types_from_files(file_paths: List[str]) -> List[str]:
    detected = set()
    for p in file_paths:
        try:
            text = extract_text_from_docx(Path(p))
        except Exception:
            text = ""
        lower = text.lower()
        for dtype, kws in DOC_TYPE_KEYWORDS.items():
            for kw in kws:
                if kw in lower:
                    detected.add(dtype)
    return list(detected)

def detect_primary_process(detected_doc_types: List[str]) -> str:
    # Simple heuristics: if doc types contain many incorporation docs -> Company Incorporation
    incorporation_set = set(PROCESS_CHECKLISTS.get("Company Incorporation", []))
    if incorporation_set.intersection(set(detected_doc_types)):
        return "Company Incorporation"
    # fallback
    return "Unknown"

def checklist_verification(process_name: str, detected_doc_types: List[str]) -> Dict:
    required = PROCESS_CHECKLISTS.get(process_name, [])
    uploaded = detected_doc_types
    missing = [d for d in required if d not in uploaded]
    return {
        "process": process_name,
        "documents_uploaded": len(uploaded),
        "required_documents": len(required),
        "missing_documents": missing
    }


# Process uploaded DOCX files (analyze + mark)
def analyze_docx_and_mark_multiple(file_paths: List[str], rag_index: RAGIndex) -> Tuple[List[str], str]:
    """
    Processes multiple .docx files:
    - Detects types across files
    - Verifies checklist
    - Runs paragraph-level RAG checks and marks flagged paragraphs in each file
    Returns list of output docx paths and a single JSON report path
    """
    # Detect doc types across files
    detected_types = detect_doc_types_from_files(file_paths)
    primary_process = detect_primary_process(detected_types)
    checklist = checklist_verification(primary_process, detected_types)

    overall_issues = []

    output_files = []
    for file_path in file_paths:
        doc = Document(file_path)
        file_issues = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            retrieved = rag_index.retrieve(text, top_k=TOP_K)
            retrieved_texts = [r[0] for r in retrieved]
            analysis = openai_analyze(text, retrieved_texts)

            if analysis.get("flag"):
                # highlight paragraph
                for run in para.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                # append inline bracketed comment with the first issue summary
                if analysis.get("issues"):
                    reason = analysis["issues"][0].get("reason", "Potential issue")
                    rec = analysis["issues"][0].get("recommendation", "")
                    para.add_run(f" [COMMENT: {reason} | Recommendation: {rec}]")
                # collect issue
                file_issues.append({
                    "document": Path(file_path).name,
                    "paragraph_text": text,
                    "analysis": analysis
                })

        # save reviewed doc
        base = Path(file_path).stem
        out_docx = f"{base}_reviewed.docx"
        doc.save(out_docx)
        output_files.append(out_docx)
        overall_issues.extend(file_issues)

    # Build JSON report
    report = {
        "process": checklist["process"],
        "documents_uploaded": checklist["documents_uploaded"],
        "required_documents": checklist["required_documents"],
        "missing_documents": checklist["missing_documents"],
        "issues_found": []
    }

    # Convert overall_issues into the required issue schema
    for it in overall_issues:
        # Attempt to detect a "section" snippet - naive: first 50 chars
        section = (it["paragraph_text"][:50] + "...") if len(it["paragraph_text"]) > 50 else it["paragraph_text"]
        reason = it["analysis"].get("issues", [{}])[0].get("reason", "Potential issue")
        rec = it["analysis"].get("issues", [{}])[0].get("recommendation", "")
        # Severity heuristic: use confidence
        conf = float(it["analysis"].get("confidence", 0.0))
        severity = "High" if conf >= 0.7 else ("Medium" if conf >= 0.4 else "Low")
        report["issues_found"].append({
            "document": it["document"],
            "section": section,
            "issue": reason,
            "severity": severity,
            "suggestion": rec
        })

    out_json = "adgm_review_report.json"
    with open(out_json, "w", encoding="utf-8") as f:
        json.dump(report, f, indent=2, ensure_ascii=False)

    return output_files, out_json


def build_index_or_exit() -> RAGIndex:
    texts = load_ref_texts(ADGM_REFS_DIR)
    if not texts:
        raise RuntimeError(f"No ADGM reference files found in {ADGM_REFS_DIR}. Place PDFs/DOCX/TXT there.")
    rag = RAGIndex()
    rag.build(texts)
    print(f"[INFO] Built index with {len(rag.texts)} chunks.")
    return rag


def launch_app():
    rag = build_index_or_exit()

    def gradio_process(files):
        # Gradio gives a list of file-like objects. Save them temporarily.
        tmp_paths = []
        for f in files:
            t = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            t.write(f.read())
            t.flush()
            tmp_paths.append(t.name)
        reviewed_files, json_report = analyze_docx_and_mark_multiple(tmp_paths, rag)
        # Return reviewed files individually (if more than one) and the JSON report
        outputs = reviewed_files + [json_report]
        # Gradio requires matching number of outputs; return as a list of downloadable file paths.
        return outputs

    # outputs: dynamic number; for simplicity allow two files max in UI (adjust as you like)
    iface = gr.Interface(
        fn=gradio_process,
        inputs=gr.Files(file_count="multiple", type="file"),
        outputs=gr.File(label="Reviewed files (download reviewed docx and JSON report)"),
        title="ADGM RAG Corporate Agent (Prototype)",
        description="Upload one or more .docx files (company formation docs). Returns marked-up DOCX(s) and a JSON report."
    )
    iface.launch()

if __name__ == "__main__":
    launch_app()
